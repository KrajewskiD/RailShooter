from pathlib import Path
import math
import zipfile

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BUILD_DIR = ROOT / "outputs" / "rozdzial_4_build"
OUT_PATH = ROOT / "outputs" / "Rozdzial_4_Implementacja_proceduralnego_generowania_terenu.docx"

FONT_SERIF = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
FONT_SERIF_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
FONT_SANS = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_SANS_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
FONT_MONO = "/System/Library/Fonts/Menlo.ttc"

INK = "#172033"
BLUE = "#24527A"
LIGHT_BLUE = "#EAF2F8"
LIGHT_GRAY = "#F2F4F7"
MID_GRAY = "#D7DEE7"
ACCENT = "#D45B3D"
GREEN = "#3E7C59"


def pil_font(path, size):
    return ImageFont.truetype(path, size=size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C1CC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        p_pr.remove(node)


def set_keep_lines(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepLines"))
    if value and node is None:
        p_pr.append(OxmlElement("w:keepLines"))
    elif not value and node is not None:
        p_pr.remove(node)


def set_run_font(run, name, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])
    set_run_font(run, "Times New Roman", 10)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for style_name, size, before, after in (
        ("Heading 1", 16, 0, 12),
        ("Heading 2", 14, 18, 8),
        ("Heading 3", 12, 14, 6),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.15

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "Rozdział 4. Implementacja proceduralnego generowania terenu"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, "Times New Roman", 9, italic=True, color="666666")

    add_page_number(section.footer.paragraphs[0])

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, "Times New Roman", 12, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, "Times New Roman", 12)
    else:
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", 12)
    return p


def add_source(doc, text, centered=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, "Times New Roman", 9, italic=True, color="555555")
    set_keep_with_next(p, False)
    return p


def add_figure(doc, image_path, caption, source, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    set_keep_with_next(p)
    cap = doc.add_paragraph(caption, style="Caption")
    set_keep_with_next(cap)
    add_source(doc, source)


def add_equation(doc, number, formula, description, variables, source):
    image_path = BUILD_DIR / f"equation_{number.replace('.', '_')}.png"
    make_equation_image(formula, image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(image_path), width=Cm(14.0))
    shape._inline.docPr.set("descr", f"Równanie {number}: {formula}")
    shape._inline.docPr.set("title", f"Równanie {number}")
    set_keep_with_next(p)

    where = doc.add_paragraph()
    where.paragraph_format.first_line_indent = Cm(0)
    where.paragraph_format.space_after = Pt(1)
    where.paragraph_format.line_spacing = 1.0
    r = where.add_run("gdzie:")
    set_run_font(r, "Times New Roman", 11, italic=True)
    set_keep_with_next(where)

    for idx, (symbol, meaning) in enumerate(variables):
        vp = doc.add_paragraph()
        vp.paragraph_format.left_indent = Cm(0.75)
        vp.paragraph_format.first_line_indent = Cm(-0.5)
        vp.paragraph_format.space_after = Pt(0)
        vp.paragraph_format.line_spacing = 1.0
        rs = vp.add_run(f"{symbol} – ")
        set_run_font(rs, "Times New Roman", 11, italic=True)
        rm = vp.add_run(meaning)
        set_run_font(rm, "Times New Roman", 11)
        set_keep_with_next(vp, idx < len(variables) - 1)

    cap = doc.add_paragraph(f"Równanie ({number}). {description}", style="Caption")
    set_keep_with_next(cap)
    add_source(doc, f"Źródło równania: {source}")


def add_code_listing(doc, number, title, code, source="opracowanie własne na podstawie kodu aplikacji"):
    lines = code.strip("\n").splitlines()
    numbered = "\n".join(f"{idx:>2}  {line}" for idx, line in enumerate(lines, 1))
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.5)
    cell = table.cell(0, 0)
    cell.width = Cm(15.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    set_table_borders(table, color="C5CDD7", size="4")
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_row_cant_split(row)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(numbered)
    set_run_font(run, "Menlo", 7.4)
    cap = doc.add_paragraph(f"Listing {number}. {title}", style="Caption")
    set_keep_with_next(cap)
    add_source(doc, f"Źródło: {source}")


def make_equation_image(text, path):
    scale = 2
    width, height = 1800, 150
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = pil_font(FONT_SERIF, 54)
    bbox = draw.textbbox((0, 0), text, font=font)
    x = (width - (bbox[2] - bbox[0])) // 2
    y = (height - (bbox[3] - bbox[1])) // 2 - 4
    draw.text((x, y), text, font=font, fill=INK)
    image.resize((width // scale, height // scale), Image.Resampling.LANCZOS).save(path, dpi=(180, 180))


def arrow(draw, start, end, color=BLUE, width=5, head=14):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    left = (
        end[0] - head * math.cos(angle - math.pi / 6),
        end[1] - head * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - head * math.cos(angle + math.pi / 6),
        end[1] - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def rounded_box(draw, xy, text, fill=LIGHT_BLUE, outline=BLUE, font=None, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x0, y0, x1, y1 = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    tx = (x0 + x1 - (bbox[2] - bbox[0])) / 2
    ty = (y0 + y1 - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((tx, ty), text, font=font, fill=INK, spacing=6, align="center")


def fade(t):
    return t * t * t * (t * (t * 6.0 - 15.0) + 10.0)


GRADS = np.array([
    [1.0, 0.0], [-1.0, 0.0], [0.0, 1.0], [0.0, -1.0],
    [0.70710678, 0.70710678], [-0.70710678, 0.70710678],
    [0.70710678, -0.70710678], [-0.70710678, -0.70710678],
])


def make_perm(seed=1337):
    rng = np.random.default_rng(seed)
    p = np.arange(256, dtype=np.int32)
    rng.shuffle(p)
    return np.concatenate([p, p])


PERM = make_perm()


def grad_dot(hash_value, x, y):
    g = GRADS[hash_value & 7]
    return g[0] * x + g[1] * y


def perlin2(x, y):
    xi = math.floor(x)
    yi = math.floor(y)
    xf = x - xi
    yf = y - yi
    x0 = xi & 255
    y0 = yi & 255
    h00 = PERM[x0 + PERM[y0]]
    h10 = PERM[x0 + 1 + PERM[y0]]
    h01 = PERM[x0 + PERM[y0 + 1]]
    h11 = PERM[x0 + 1 + PERM[y0 + 1]]
    n00 = grad_dot(h00, xf, yf)
    n10 = grad_dot(h10, xf - 1.0, yf)
    n01 = grad_dot(h01, xf, yf - 1.0)
    n11 = grad_dot(h11, xf - 1.0, yf - 1.0)
    u, v = fade(xf), fade(yf)
    nx0 = n00 + (n10 - n00) * u
    nx1 = n01 + (n11 - n01) * u
    return (nx0 + (nx1 - nx0) * v) * 1.45


def simplex2(x, y):
    f2 = (math.sqrt(3.0) - 1.0) / 2.0
    g2 = (3.0 - math.sqrt(3.0)) / 6.0
    skew = (x + y) * f2
    i = math.floor(x + skew)
    j = math.floor(y + skew)
    unskew = (i + j) * g2
    x0 = x - (i - unskew)
    y0 = y - (j - unskew)
    if x0 > y0:
        i1, j1 = 1, 0
    else:
        i1, j1 = 0, 1
    x1, y1 = x0 - i1 + g2, y0 - j1 + g2
    x2, y2 = x0 - 1.0 + 2.0 * g2, y0 - 1.0 + 2.0 * g2
    ii, jj = i & 255, j & 255
    hashes = (
        PERM[ii + PERM[jj]],
        PERM[ii + i1 + PERM[jj + j1]],
        PERM[ii + 1 + PERM[jj + 1]],
    )
    value = 0.0
    for h, dx, dy in zip(hashes, (x0, x1, x2), (y0, y1, y2)):
        t = 0.5 - dx * dx - dy * dy
        if t > 0:
            t *= t
            value += t * t * grad_dot(h, dx, dy)
    return max(-1.0, min(1.0, value * 70.0))


def feature_point(cx, cy):
    h = int(PERM[(cx & 255) + PERM[cy & 255]])
    h2 = int(PERM[(h + 71) & 255])
    return cx + 0.15 + 0.7 * (h / 255.0), cy + 0.15 + 0.7 * (h2 / 255.0)


def worley2(x, y):
    xi, yi = math.floor(x), math.floor(y)
    best = 1e9
    for dy in (-1, 0, 1):
        for dx in (-1, 0, 1):
            px, py = feature_point(xi + dx, yi + dy)
            dist = math.sqrt((px - x) ** 2 + (py - y) ** 2)
            best = min(best, dist)
    return best


def sample_field(func, size=168, domain=6.0):
    arr = np.empty((size, size), dtype=np.float32)
    for iy in range(size):
        y = iy / (size - 1) * domain
        for ix in range(size):
            x = ix / (size - 1) * domain
            arr[iy, ix] = func(x, y)
    return arr


def normalize_image(arr, symmetric=False):
    if symmetric:
        arr = np.clip((arr + 1.0) * 0.5, 0.0, 1.0)
    else:
        lo, hi = float(arr.min()), float(arr.max())
        arr = (arr - lo) / max(1e-9, hi - lo)
    return Image.fromarray((arr * 255.0).astype(np.uint8), mode="L").convert("RGB")


def make_noise_comparison(path):
    panels = [
        ("Perlin", normalize_image(sample_field(perlin2), symmetric=True)),
        ("Simplex", normalize_image(sample_field(simplex2), symmetric=True)),
        ("Worley (F1)", normalize_image(sample_field(worley2), symmetric=False)),
    ]
    canvas = Image.new("RGB", (1740, 650), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 35)
    label_font = pil_font(FONT_SANS, 25)
    draw.text((60, 34), "Porównanie szumów przy wspólnym obszarze próbkowania", font=title_font, fill=INK)
    for idx, (label, img) in enumerate(panels):
        x = 60 + idx * 560
        scaled = img.resize((500, 500), Image.Resampling.LANCZOS)
        canvas.paste(scaled, (x, 100))
        draw.rectangle((x, 100, x + 500, 600), outline=MID_GRAY, width=3)
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.rectangle((x, 555, x + 500, 600), fill=(255, 255, 255, 235))
        draw.text((x + 250 - (bbox[2] - bbox[0]) / 2, 563), label, font=label_font, fill=INK)
    canvas.save(path, dpi=(180, 180))


def make_simplex_geometry(path):
    canvas = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 34)
    label_font = pil_font(FONT_SANS, 24)
    small_font = pil_font(FONT_SANS, 20)
    draw.text((60, 40), "Próbkowanie dwuwymiarowej komórki Simplex", font=title_font, fill=INK)

    origin_x, origin_y = 190, 730
    sx, sy = 160, 138
    points = {}
    for j in range(-1, 6):
        for i in range(-1, 7):
            x = origin_x + (i + 0.5 * j) * sx
            y = origin_y - j * sy
            points[(i, j)] = (x, y)
    for j in range(-1, 5):
        for i in range(-1, 6):
            p = points[(i, j)]
            for q in ((i + 1, j), (i, j + 1), (i + 1, j - 1)):
                if q in points:
                    draw.line([p, points[q]], fill="#C9D2DC", width=3)

    a = points[(2, 2)]
    b = points[(3, 2)]
    c = points[(2, 3)]
    draw.polygon([a, b, c], fill="#DDEBF5", outline=BLUE)
    draw.line([a, b], fill=BLUE, width=6)
    draw.line([b, c], fill=BLUE, width=6)
    draw.line([c, a], fill=BLUE, width=6)

    px = (a[0] * 0.28 + b[0] * 0.32 + c[0] * 0.40)
    py = (a[1] * 0.28 + b[1] * 0.32 + c[1] * 0.40)
    for label, corner in (("p0", a), ("p1", b), ("p2", c)):
        draw.ellipse((corner[0] - 10, corner[1] - 10, corner[0] + 10, corner[1] + 10), fill=BLUE)
        draw.text((corner[0] + 14, corner[1] - 28), label, font=label_font, fill=INK)
        arrow(draw, (px, py), corner, color=GREEN, width=4, head=12)
    draw.ellipse((px - 12, py - 12, px + 12, py + 12), fill=ACCENT)
    draw.text((px + 16, py - 38), "punkt p", font=label_font, fill=ACCENT)

    note_x = 980
    rounded_box(draw, (note_x, 200, 1400, 340), "1. Skoś układu\nwspółrzędnych", font=small_font)
    rounded_box(draw, (note_x, 390, 1400, 530), "2. Wybór trzech\nwierzchołków", font=small_font)
    rounded_box(draw, (note_x, 580, 1400, 720), "3. Suma tłumionych\niloczynów skalarnych", font=small_font)
    arrow(draw, (1190, 340), (1190, 390), color=BLUE)
    arrow(draw, (1190, 530), (1190, 580), color=BLUE)
    draw.text((75, 820), "W 2D komórka simpleksowa jest trójkątem, dlatego wynik tworzą trzy wkłady narożne.", font=small_font, fill="#555555")
    canvas.save(path, dpi=(180, 180))


def make_fractal_comparison(path):
    size = 168
    base = sample_field(lambda x, y: perlin2(x * 0.7, y * 0.7), size=size, domain=6.0)
    fbm = np.zeros_like(base)
    ridge = np.zeros_like(base)
    amp = 1.0
    amp_sum = 0.0
    for octave in range(5):
        freq = 0.7 * (2.0 ** octave)
        layer = sample_field(lambda x, y, f=freq: perlin2(x * f, y * f), size=size, domain=6.0)
        fbm += layer * amp
        ridge += (1.0 - 2.0 * np.abs(layer)) * amp
        amp_sum += amp
        amp *= 0.5
    fbm /= amp_sum
    ridge /= amp_sum
    panels = [
        ("Jedna oktawa", normalize_image(base, symmetric=True)),
        ("fBM, 5 oktaw", normalize_image(fbm, symmetric=True)),
        ("Ridged, 5 oktaw", normalize_image(ridge, symmetric=True)),
    ]
    canvas = Image.new("RGB", (1740, 650), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 35)
    label_font = pil_font(FONT_SANS, 25)
    draw.text((60, 34), "Wpływ syntezy wielooktawowej na strukturę sygnału", font=title_font, fill=INK)
    for idx, (label, img) in enumerate(panels):
        x = 60 + idx * 560
        canvas.paste(img.resize((500, 500), Image.Resampling.LANCZOS), (x, 100))
        draw.rectangle((x, 100, x + 500, 600), outline=MID_GRAY, width=3)
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.rectangle((x, 555, x + 500, 600), fill="white")
        draw.text((x + 250 - (bbox[2] - bbox[0]) / 2, 563), label, font=label_font, fill=INK)
    canvas.save(path, dpi=(180, 180))


def make_pipeline(path):
    canvas = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 34)
    box_font = pil_font(FONT_SANS, 21)
    draw.text((55, 38), "Przepływ danych podczas generowania fragmentu terenu", font=title_font, fill=INK)
    labels = [
        "Współrzędne\nświata",
        "Warstwy\nszumu",
        "fBM lub\nRidged",
        "Mieszanie\nwarstw",
        "Krzywa\nwysokości",
        "Klimat i\nbiom",
        "Wierzchołki,\nnormale, indeksy",
    ]
    x = 55
    boxes = []
    for idx, label in enumerate(labels):
        width = 210 if idx < 6 else 270
        box = (x, 240, x + width, 415)
        fill = LIGHT_BLUE if idx < 5 else ("#E9F4EC" if idx == 5 else "#FFF0E9")
        outline = BLUE if idx < 5 else (GREEN if idx == 5 else ACCENT)
        rounded_box(draw, box, label, fill=fill, outline=outline, font=box_font)
        boxes.append(box)
        x += width + 35
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 3, (left[1] + left[3]) / 2), (right[0] - 3, (right[1] + right[3]) / 2), color="#65758A", width=4, head=12)
    draw.text((70, 505), "Wynik próbkowania jest najpierw kształtowany w dziedzinie wysokości, a dopiero później używany do budowy geometrii i koloru.", font=box_font, fill="#555555")
    canvas.save(path, dpi=(180, 180))


def make_job_graph(path):
    canvas = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 34)
    box_font = pil_font(FONT_SANS, 23)
    small_font = pil_font(FONT_SANS, 20)
    draw.text((55, 38), "Łańcuch zależności zadań generowania siatki", font=title_font, fill=INK)
    boxes = [
        (120, 210, 500, 360, "Próbkowanie wysokości,\nklimatu i koloru"),
        (560, 210, 940, 360, "Obliczanie normalnych\nz sąsiednich wysokości"),
        (1000, 210, 1380, 360, "Budowa indeksów\ntrójkątów"),
    ]
    for idx, box in enumerate(boxes):
        rounded_box(draw, box[:4], box[4], fill=LIGHT_BLUE if idx < 2 else "#FFF0E9", outline=BLUE if idx < 2 else ACCENT, font=box_font)
    arrow(draw, (500, 285), (560, 285), color="#65758A")
    arrow(draw, (940, 285), (1000, 285), color="#65758A")
    rounded_box(draw, (500, 520, 1000, 675), "Sprawdzenie JobHandle.IsCompleted\noraz przekazanie wyniku do kolejki", fill="#E9F4EC", outline=GREEN, font=box_font)
    arrow(draw, (1190, 360), (880, 520), color=GREEN)
    draw.text((215, 420), "IJobParallelFor", font=small_font, fill="#555555")
    draw.text((655, 420), "zależność: handle1", font=small_font, fill="#555555")
    draw.text((1090, 420), "zależność: handle2", font=small_font, fill="#555555")
    draw.text((380, 760), "Modyfikacja obiektu Mesh pozostaje operacją wykonywaną na wątku głównym.", font=small_font, fill="#555555")
    canvas.save(path, dpi=(180, 180))


def make_biome_matrix(path):
    canvas = Image.new("RGB", (1200, 940), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(FONT_SANS_BOLD, 34)
    label_font = pil_font(FONT_SANS, 24)
    cell_font = pil_font(FONT_SANS_BOLD, 20)
    draw.text((55, 35), "Macierz barw zależna od temperatury i wilgotności", font=title_font, fill=INK)
    colors = [
        [(215, 225, 240), (165, 215, 175), (45, 135, 80)],
        [(245, 215, 110), (170, 230, 90), (55, 165, 80)],
        [(255, 175, 90), (220, 200, 70), (35, 150, 50)],
    ]
    labels = [
        ["zimny i suchy", "zimny", "zimny i wilgotny"],
        ["umiarkowany i suchy", "umiarkowany", "umiarkowany i wilgotny"],
        ["gorący i suchy", "gorący", "gorący i wilgotny"],
    ]
    x0, y0, cell_w, cell_h = 250, 170, 285, 210
    for row in range(3):
        for col in range(3):
            x = x0 + col * cell_w
            y = y0 + row * cell_h
            draw.rectangle((x, y, x + cell_w, y + cell_h), fill=colors[row][col], outline="white", width=5)
            text = labels[row][col]
            bbox = draw.multiline_textbbox((0, 0), text, font=cell_font, spacing=4, align="center")
            draw.multiline_text(
                (x + cell_w / 2 - (bbox[2] - bbox[0]) / 2, y + cell_h / 2 - (bbox[3] - bbox[1]) / 2),
                text,
                font=cell_font,
                fill=INK,
                spacing=4,
                align="center",
            )
    draw.text((420, 825), "wilgotność  →", font=label_font, fill=INK)
    draw.text((40, 510), "temperatura\n↓", font=label_font, fill=INK, spacing=4)
    draw.text((x0 + 15, 135), "sucho", font=label_font, fill="#555555")
    draw.text((x0 + cell_w + 95, 135), "średnio", font=label_font, fill="#555555")
    draw.text((x0 + 2 * cell_w + 180, 135), "wilgotno", font=label_font, fill="#555555")
    canvas.save(path, dpi=(180, 180))


def generate_figures():
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "noise": BUILD_DIR / "noise_comparison.png",
        "simplex": BUILD_DIR / "simplex_geometry.png",
        "fractal": BUILD_DIR / "fractal_comparison.png",
        "pipeline": BUILD_DIR / "terrain_pipeline.png",
        "jobs": BUILD_DIR / "job_dependencies.png",
        "biomes": BUILD_DIR / "biome_matrix.png",
    }
    make_noise_comparison(figures["noise"])
    make_simplex_geometry(figures["simplex"])
    make_fractal_comparison(figures["fractal"])
    make_pipeline(figures["pipeline"])
    make_job_graph(figures["jobs"])
    make_biome_matrix(figures["biomes"])
    return figures


CONFIG_LISTING = """
public enum NoiseType { Perlin = 0, SimplexNoise = 1, Worley = 2 }
public enum FractalType { None, FBm, Ridged }
public enum BaseNoiseBlendMode { Add, Subtract, Multiply, Max, Min }

[System.Serializable]
public class BaseNoiseLayer
{
    public bool enabled = true;
    public BaseNoiseBlendMode blendMode = BaseNoiseBlendMode.Add;
    [Range(0f, 2f)] public float strength = 1f;
    public NoiseType type = NoiseType.SimplexNoise;
    public FractalType fractal = FractalType.FBm;
    public float frequency = 0.002f;
    [Range(1, 8)] public int octaves = 2;
    public float lacunarity = 2f;
    public float persistence = 0.5f;
    public CellularDistance cellDistance = CellularDistance.EuclideanSq;
    public CellularReturn cellReturn = CellularReturn.Distance;
    [Range(0f, 1f)] public float jitter = 1f;
}
"""

PERM_LISTING = """
public static void BuildPermutation512(int seed, NativeArray<int> outPerm)
{
    for (int i = 0; i < 256; i++)
        outPerm[i] = BasePermutation256[i];

    uint state = seed == 0 ? 0x9E3779B9u : (uint)seed;
    for (int i = 255; i > 0; i--)
    {
        state = XorShift32(state);
        int j = (int)(state % (uint)(i + 1));
        (outPerm[i], outPerm[j]) = (outPerm[j], outPerm[i]);
    }

    for (int i = 256; i < 512; i++)
        outPerm[i] = outPerm[i - 256];
}

private static uint XorShift32(uint x)
{
    x ^= x << 13;
    x ^= x >> 17;
    x ^= x << 5;
    return x;
}
"""

PERLIN_LISTING = """
private static float Perlin2DPerm(NativeArray<int> perm, float x, float y)
{
    int xi = FastFloor(x), yi = FastFloor(y);
    float xd0 = x - xi, yd0 = y - yi;
    float xd1 = xd0 - 1f, yd1 = yd0 - 1f;
    float xs = Fade(xd0), ys = Fade(yd0);

    int xi0 = xi & 255;
    int yi0 = yi & 255;
    int h00 = perm[xi0 + perm[yi0]];
    int h10 = perm[xi0 + 1 + perm[yi0]];
    int h01 = perm[xi0 + perm[yi0 + 1]];
    int h11 = perm[xi0 + 1 + perm[yi0 + 1]];

    float n00 = Grad(h00, xd0, yd0);
    float n10 = Grad(h10, xd1, yd0);
    float n01 = Grad(h01, xd0, yd1);
    float n11 = Grad(h11, xd1, yd1);

    return Lerp(Lerp(n00, n10, xs),
                Lerp(n01, n11, xs), ys) * 1.4247691f;
}
"""

SIMPLEX_LISTING = """
private static float Noise2DSimplexPermutation(
    float x, float y, NativeArray<int> perm)
{
    const float F2 = 0.3660254037844386f;
    const float G2 = 0.2113248654051871f;

    float skew = (x + y) * F2;
    int i = FastFloor(x + skew);
    int j = FastFloor(y + skew);

    float unskew = (i + j) * G2;
    float x0 = x - (i - unskew);
    float y0 = y - (j - unskew);

    int i1 = x0 > y0 ? 1 : 0;
    int j1 = x0 > y0 ? 0 : 1;
    float x1 = x0 - i1 + G2;
    float y1 = y0 - j1 + G2;
    float x2 = x0 - 1f + 2f * G2;
    float y2 = y0 - 1f + 2f * G2;

    int ii = i & 255;
    int jj = j & 255;
    int g0 = perm[ii + perm[jj]];
    int g1 = perm[ii + i1 + perm[jj + j1]];
    int g2 = perm[ii + 1 + perm[jj + 1]];

    float value = Kernel2DPerm(g0, x0, y0)
                + Kernel2DPerm(g1, x1, y1)
                + Kernel2DPerm(g2, x2, y2);
    return math.clamp(value * 70f, -1f, 1f);
}
"""

FBM_LISTING = """
private float FBm2D(float x, float y, NativeArray<int> permutation)
{
    float sum = 0f, amp = 1f, maxAmp = 0f;
    int s = seed;
    int oct = octaves < 1 ? 1 : octaves;

    for (int i = 0; i < oct; i++)
    {
        sum += Single2D(s, x, y, permutation) * amp;
        maxAmp += amp;
        amp *= gain;
        x *= lacunarity;
        y *= lacunarity;
        s = unchecked(s + 0x68E31DA4);
    }

    return sum / maxAmp;
}
"""

RIDGED_LISTING = """
private float Ridged2D(float x, float y, NativeArray<int> permutation)
{
    float sum = 0f, amp = 1f, maxAmp = 0f;
    int s = seed;
    int oct = octaves < 1 ? 1 : octaves;

    for (int i = 0; i < oct; i++)
    {
        float n = Single2D(s, x, y, permutation);
        float absN = n < 0f ? -n : n;
        sum += (1f - 2f * absN) * amp;
        maxAmp += amp;
        amp *= gain;
        x *= lacunarity;
        y *= lacunarity;
        s = unchecked(s + 0x68E31DA4);
    }

    return sum / maxAmp;
}
"""

BLEND_LISTING = """
for (int i = 0; i < count; i++)
{
    BaseNoiseLayerRuntime layer = layers[i];
    if (layer.enabled == 0 || layer.strength <= 0f)
        continue;

    float sample = noises[i].GetNoise2D(x, z, permutation);
    float weighted = sample * layer.strength;

    if (!hasValue)
    {
        value = weighted;
        hasValue = true;
        continue;
    }

    switch (layer.blendMode)
    {
        case (int)BaseNoiseBlendMode.Subtract:
            value -= weighted;
            break;
        case (int)BaseNoiseBlendMode.Multiply:
            value *= math.max(0f, 1f + weighted);
            break;
        case (int)BaseNoiseBlendMode.Max:
            value = math.max(value, weighted);
            break;
        case (int)BaseNoiseBlendMode.Min:
            value = math.min(value, weighted);
            break;
        default:
            value += weighted;
            break;
    }
}
return hasValue ? math.clamp(value, -1f, 1f) : 0f;
"""

TERRAIN_EXECUTE_LISTING = """
int x = index % borderedVCount;
int y = index / borderedVCount;
float worldX = startX + (x - 1) * resolution;
float worldZ = startZ + (y - 1) * resolution;

float baseHeight = NoiseProvider.SampleBaseNoiseStack(
    baseNoises, baseLayerSettings, permutation, worldX, worldZ);

float t = math.saturate((baseHeight + 1f) * 0.5f);
float idx = t * (heightCurveLUT.Length - 1);
int i0 = (int)idx;
int i1 = math.min(i0 + 1, heightCurveLUT.Length - 1);
float shaped = math.lerp(
    heightCurveLUT[i0], heightCurveLUT[i1], idx - i0);
float finalHeight = math.max(0f, shaped * fAmp);

ClimateSample climate = SampleClimate(
    temperatureNoise, moistureNoise, permutation,
    worldX, worldZ, finalHeight, biomeParams, true);

chunkData[index] = new float4(
    finalHeight, climate.temperature, climate.moisture, decoMask);
"""

NORMAL_LISTING = """
int mx = index % vCount;
int my = index / vCount;
int hx = mx + 1;
int hy = my + 1;

float hL = chunkData[hy * borderedVCount + hx - 1].x;
float hR = chunkData[hy * borderedVCount + hx + 1].x;
float hD = chunkData[(hy - 1) * borderedVCount + hx].x;
float hU = chunkData[(hy + 1) * borderedVCount + hx].x;

var v = vertexData[index];
v.normal = math.normalize(new float3(
    hL - hR, 2f * resolution, hD - hU));
vertexData[index] = v;
"""

SCHEDULE_LISTING = """
JobHandle handle1 =
    terrainJob.Schedule(totalBorderedVertices, 64);

var normalsJob = new CalculateNormalsJob
{
    chunkData = chunkData,
    vertexData = vertexBuffer,
    borderedVCount = borderedVerticesPerLine,
    vCount = verticesPerLine,
    resolution = resolution
};
JobHandle handle2 =
    normalsJob.Schedule(totalVertices, 64, handle1);

JobHandle handle3 =
    trianglesJob.Schedule(segments * segments, 64, handle2);

meshHandle = handle3;
state = GenState.Meshing;
"""

CLIMATE_LISTING = """
public static float NormalizeTemperatureNoise(float rawNoise)
{
    rawNoise = math.sign(rawNoise) * math.sqrt(math.abs(rawNoise));
    return math.saturate((rawNoise + 1f) * 0.5f);
}

public static float NormalizeMoistureNoise(float rawNoise)
{
    return math.saturate((rawNoise + 1f) * 0.5f);
}

float peakHeight = math.max(0.001f, biomeParams.peakHeight);
float height01 = math.saturate(terrainHeight / peakHeight);
temperature = math.saturate(
    temperature - height01 * biomeParams.temperatureAltitudeCooling);

float valleyTop = math.max(0.001f, biomeParams.valleyTop);
float valleyWetness =
    (1f - math.saturate(terrainHeight / valleyTop))
    * biomeParams.valleyMoistureBoost;
float peakDryness = height01 * biomeParams.peakMoistureDryness;
moisture = math.saturate(moisture + valleyWetness);
moisture *= math.saturate(1f - peakDryness);
"""

COLOR_LISTING = """
if (height < p.seaLevel)
    return Sea;

if (height < p.beachTop)
    return Color32Lerp(
        Sea, Beach, math.smoothstep(p.seaLevel, p.beachTop, height));

Color32 biome = GetZoneColor(
    temperature, moisture,
    ColdDry, ColdMid, ColdWet,
    NormDry, NormMid, NormWet,
    HotDry, HotMid, HotWet);

if (height < p.biomeTop)
    return Color32Lerp(
        Beach, biome, math.smoothstep(p.beachTop, p.biomeTop, height));

float rockBlend =
    math.smoothstep(p.rockStart, p.rockEnd, height) * 0.5f;
biome = Color32Lerp(biome, Rock, rockBlend);

float snowLine = math.lerp(
    p.snowLineCold, p.snowLineHot, math.saturate(temperature));
float snowBlend =
    math.smoothstep(snowLine, snowLine + p.snowBandWidth, height);
return Color32Lerp(biome, Snow, snowBlend);
"""


def build_document(figures):
    doc = Document()
    style_document(doc)

    title = add_heading(doc, "4. Implementacja proceduralnego generowania terenu", 1)
    title.paragraph_format.space_after = Pt(12)

    add_body(
        doc,
        "Proceduralne generowanie terenu stanowi jeden z głównych elementów aplikacji. "
        "Proces ten obejmuje utworzenie mapy wysokości, zbudowanie siatki fragmentów terenu, "
        "wyznaczenie parametrów klimatycznych oraz przypisanie kolorów powierzchni. "
        "Poszczególne etapy wykonywane są dynamicznie w zależności od położenia gracza. "
        "Zastosowanie systemu zadań Unity pozwala przenieść najbardziej kosztowne obliczenia "
        "poza główny wątek aplikacji."
    )
    add_body(
        doc,
        "W niniejszym rozdziale algorytmy szumu opisano na podstawie ich publikowanych definicji, "
        "a decyzje charakterystyczne dla aplikacji, takie jak kolejność mieszania warstw, kształtowanie "
        "mapy wysokości i reguły kolorowania, przedstawiono oddzielnie jako rozwiązania projektowe. "
        "Takie rozdzielenie umożliwia wskazanie, które elementy wynikają bezpośrednio z literatury, "
        "a które odpowiadają za sposób wykorzystania szumu w generatorze terenu."
    )

    add_heading(doc, "4.1. Konfiguracja i tworzenie szumów", 2)
    add_body(
        doc,
        "Podstawą zaimplementowanego systemu generowania terenu są funkcje szumu, które dla "
        "określonych współrzędnych świata zwracają wartości wykorzystywane do wyznaczenia wysokości, "
        "temperatury lub wilgotności. Każdy z zastosowanych algorytmów charakteryzuje się innym sposobem "
        "obliczania wartości. Szum Perlina należy do grupy szumów gradientowych [1, 2], Simplex rozwija "
        "tę ideę z użyciem komórek simpleksowych [3, 4], natomiast szum Worleya bazuje na odległościach "
        "od rozproszonych punktów cech [5]."
    )
    add_body(
        doc,
        "Parametry pojedynczego źródła szumu powinny umożliwiać zarówno wybór funkcji bazowej, jak i "
        "określenie jej skali, liczby oktaw oraz sposobu połączenia z pozostałymi źródłami. "
        "Dla algorytmu komórkowego potrzebne są ponadto ustawienia metryki odległości, zwracanej "
        "wartości oraz amplitudy przesunięcia punktu cechy wewnątrz komórki."
    )
    add_code_listing(doc, "4.14", "Struktura konfiguracji pojedynczej warstwy szumu", CONFIG_LISTING)
    add_body(
        doc,
        "Fragment kodu zamieszczony na listingu 4.14 definiuje kompletny zestaw danych konfiguracyjnych "
        "warstwy. Pole type wybiera algorytm bazowy, frequency określa częstotliwość próbkowania, "
        "a strength skaluje udział próbki w wyniku. Pola octaves, lacunarity i persistence są używane "
        "przez warianty wielooktawowe. Ustawienia cellDistance, cellReturn i jitter dotyczą wyłącznie "
        "szumu Worleya. Dzięki temu wszystkie źródła mogą być przechowywane w jednej uporządkowanej "
        "kolekcji, mimo że część parametrów jest interpretowana tylko przez wybrane algorytmy."
    )

    add_body(
        doc,
        "Porównanie kształtu funkcji bazowych wymaga zachowania jednakowego obszaru współrzędnych, "
        "częstotliwości i rozdzielczości obrazu. Zmiana któregokolwiek z tych parametrów mogłaby zostać "
        "błędnie uznana za cechę samego algorytmu."
    )
    add_figure(
        doc,
        figures["noise"],
        "Rysunek 4.1. Porównanie funkcji Perlin, Simplex i Worley dla wspólnego obszaru próbkowania",
        "Źródło: opracowanie własne na podstawie [2], [4] i [5].",
    )
    add_body(
        doc,
        "Na rysunku 4.1 przedstawiono trzy mapy w skali szarości. Perlin i Simplex tworzą ciągłe pola "
        "gradientowe, ale różnią się geometrią komórek i sposobem sumowania wkładów. Worley ujawnia "
        "natomiast granice obszarów związanych z najbliższymi punktami cech. Z tego powodu nie należy "
        "oczekiwać, że wszystkie trzy funkcje przy jednakowej częstotliwości wygenerują podobny wzór."
    )

    add_heading(doc, "4.1.1. Tablica permutacji i wybór gradientów", 3)
    add_body(
        doc,
        "Funkcje gradientowe muszą w sposób deterministyczny przypisać wektor gradientu do całkowitych "
        "współrzędnych komórki. Klasyczne implementacje realizują to zadanie za pomocą tablicy permutacji, "
        "która pełni rolę krótkookresowej funkcji mieszającej indeksy [1, 2, 8]. Wariant ten jest szczególnie "
        "przydatny w opisie implementacyjnym, ponieważ odpowiada konstrukcji publikowanej dla szumu Perlina "
        "i czytelnie oddziela wybór gradientu od interpolacji."
    )
    add_body(
        doc,
        "Tablica używana podczas pracy zadań ma długość 512, ale zawiera dwukrotnie powtórzoną permutację "
        "liczb od 0 do 255. Przed jej powieleniem kolejność elementów jest deterministycznie zmieniana na "
        "podstawie ziarna. Do tego celu można połączyć algorytm tasowania Fishera-Yatesa z generatorem "
        "Xorshift opisanym przez Marsaglię [9, 10]."
    )
    add_code_listing(
        doc,
        "4.15",
        "Tworzenie tablicy permutacji zależnej od ziarna",
        PERM_LISTING,
        "opracowanie własne na podstawie [8], [9] i [10]",
    )
    add_body(
        doc,
        "Na listingu 4.15 najpierw kopiowana jest permutacja bazowa, następnie każdy element od końca "
        "tablicy zamieniany jest z elementem o indeksie należącym do aktualnie nieprzetworzonego zakresu. "
        "Wartość pseudolosowa jest wyznaczana przez trzy operacje przesunięcia bitowego i XOR. Ostatnia "
        "pętla powiela pierwsze 256 pozycji, dzięki czemu zagnieżdżone indeksowanie nie wymaga dodatkowej "
        "operacji modulo. Należy odróżnić tę konstrukcję od arytmetycznego haszowania współrzędnych. "
        "Oba podejścia mogą być przedmiotem pomiaru wydajności, lecz opisane listingi szumów korzystają "
        "z tablicy permutacji, aby zachować bezpośrednią zgodność z literaturą."
    )

    add_heading(doc, "4.1.2. Szum Perlina", 3)
    add_body(
        doc,
        "Szum Perlina jest funkcją gradientową. W każdym narożniku komórki regularnej wybierany jest "
        "pseudolosowy gradient, a następnie obliczany jest jego iloczyn skalarny z wektorem prowadzącym "
        "od narożnika do punktu próbkowania [1]. Wersja poprawiona zastępuje wcześniejszą funkcję "
        "wygładzającą wielomianem piątego stopnia, którego pierwsza i druga pochodna zanikają na końcach "
        "przedziału [2]."
    )
    add_body(
        doc,
        "Wartość parametru interpolacji musi zmieniać się od zera do jedności bez skoków pochodnych na "
        "granicach sąsiednich komórek. Warunek ten realizuje wielomian stosowany przed interpolacją "
        "wkładów narożnych."
    )
    add_equation(
        doc,
        "4.1",
        "f(t) = 6t^5 - 15t^4 + 10t^3",
        "Funkcja wygładzająca stosowana w poprawionym szumie Perlina",
        [
            ("f(t)", "wartość funkcji wygładzającej"),
            ("t", "lokalna współrzędna interpolacji należąca do przedziału [0, 1]"),
        ],
        "K. Perlin, „Improving Noise”, SIGGRAPH 2002, DOI: 10.1145/566570.566636 [2].",
    )
    add_body(
        doc,
        "Równanie (4.1) służy do wyznaczenia wag interpolacji wzdłuż obu osi. Dla t = 0 i t = 1 funkcja "
        "przyjmuje odpowiednio wartości 0 i 1, natomiast zerowe pochodne na końcach przedziału ograniczają "
        "widoczne przejścia między komórkami."
    )
    add_body(
        doc,
        "Po wyznaczeniu czterech indeksów gradientów obliczane są wkłady narożne, a następnie dwie "
        "interpolacje wzdłuż osi x i jedna interpolacja wzdłuż osi y. Taka kolejność odpowiada "
        "dwuliniowemu łączeniu wartości w komórce."
    )
    add_code_listing(
        doc,
        "4.16",
        "Próbkowanie dwuwymiarowego szumu Perlina z tablicą permutacji",
        PERLIN_LISTING,
        "opracowanie własne na podstawie [1] i [2]",
    )
    add_body(
        doc,
        "Na listingu 4.16 współrzędne xi i yi wskazują komórkę zawierającą punkt, natomiast pary xd i yd "
        "opisują wektory od czterech narożników. Funkcja Fade realizuje równanie (4.1). Zagnieżdżone "
        "odczyty z tablicy permutacji wyznaczają indeksy h00, h10, h01 i h11, a funkcja Grad oblicza "
        "iloczyny skalarne. Ostatnia instrukcja wykonuje interpolację w dwóch wymiarach i skaluje wynik "
        "do zakresu używanego przez dalsze etapy generatora."
    )

    add_heading(doc, "4.1.3. Szum Simplex", 3)
    add_body(
        doc,
        "Szum Simplex został zaproponowany jako rozwinięcie klasycznego szumu gradientowego, które "
        "zastępuje hipersześcienne komórki simpleksami [3]. W dwóch wymiarach simpleksem jest trójkąt, "
        "dlatego wartość próbki zależy od trzech, a nie czterech narożników. Algorytm nie wykonuje "
        "sekwencyjnej interpolacji osiowej. Zamiast tego sumuje radialnie tłumione wkłady narożne [4]."
    )
    add_body(
        doc,
        "Regularna siatka kwadratowa jest najpierw poddawana transformacji skośnej. Pozwala to wskazać "
        "komórkę w prostym układzie całkowitym, po czym transformacja odwrotna przywraca położenie jej "
        "początku w przestrzeni wejściowej."
    )
    add_equation(
        doc,
        "4.2",
        "F_2 = (sqrt(3) - 1) / 2,     G_2 = (3 - sqrt(3)) / 6",
        "Stałe transformacji skośnej i odwrotnej dla Simplex 2D",
        [
            ("F₂", "współczynnik transformacji skośnej przestrzeni wejściowej"),
            ("G₂", "współczynnik transformacji odwrotnej"),
            ("√3", "pierwiastek kwadratowy z liczby 3 wynikający z geometrii trójkąta równobocznego"),
        ],
        "S. Gustavson, „Simplex noise demystified”, Linköping University, 2005 [4].",
    )
    add_body(
        doc,
        "W równaniu (4.2) stała F₂ jest używana do obliczenia wspólnego przesunięcia obu współrzędnych. "
        "Po zaokrągleniu skośnych współrzędnych do liczb całkowitych stała G₂ umożliwia wyznaczenie "
        "położenia początku komórki w pierwotnej przestrzeni."
    )
    add_body(
        doc,
        "W obrębie znalezionej komórki porównanie lokalnych współrzędnych określa, który z dwóch "
        "trójkątów zawiera punkt. Wybór ten ustala drugi narożnik simpleksu; pierwszy i trzeci narożnik "
        "pozostają wspólne dla obu przypadków."
    )
    add_figure(
        doc,
        figures["simplex"],
        "Rysunek 4.2. Geometria próbkowania dwuwymiarowej komórki Simplex",
        "Źródło: opracowanie własne na podstawie [3] i [4].",
    )
    add_body(
        doc,
        "Na rysunku 4.2 zaznaczono punkt p oraz trzy wierzchołki komórki, które uczestniczą w obliczeniu "
        "wyniku. Zielone odcinki odpowiadają wektorom używanym w iloczynach skalarnych. Pozostałe "
        "wierzchołki siatki nie wnoszą wkładu do danej próbki, co ogranicza liczbę obliczeń."
    )
    add_body(
        doc,
        "Każdy narożnik wpływa na wynik jedynie wewnątrz ograniczonego promienia. Wkład jest iloczynem "
        "wartości gradientowej oraz czwartej potęgi funkcji tłumiącej zależnej od kwadratu odległości."
    )
    add_equation(
        doc,
        "4.3",
        "n_i = max(0, 0.5 - ||x_i||^2)^4 * (g_i . x_i)",
        "Wkład pojedynczego narożnika w dwuwymiarowym szumie Simplex",
        [
            ("nᵢ", "wkład i-tego narożnika"),
            ("xᵢ", "wektor od i-tego narożnika do punktu próbkowania"),
            ("gᵢ", "gradient przypisany do i-tego narożnika"),
            ("‖xᵢ‖²", "kwadrat odległości punktu od narożnika"),
            ("gᵢ · xᵢ", "iloczyn skalarny gradientu i wektora lokalnego"),
        ],
        "S. Gustavson, „Simplex noise demystified”, Linköping University, 2005 [4].",
    )
    add_body(
        doc,
        "Równanie (4.3) powoduje zanik wkładu poza obszarem wpływu narożnika. Wewnątrz tego obszaru "
        "czwarta potęga zapewnia gładkie tłumienie, a znak i kierunek zmiany wartości wynikają z "
        "iloczynu skalarnego gradientu z wektorem lokalnym. Końcowa próbka jest sumą trzech takich "
        "wkładów pomnożoną przez stałą normalizującą."
    )
    add_body(
        doc,
        "Implementacja dwuwymiarowa może zostać zapisana bez rozgałęzionej procedury sortowania osi. "
        "Wystarcza porównanie x0 i y0, ponieważ w dwóch wymiarach istnieją tylko dwa możliwe porządki "
        "przejścia przez trójkąt."
    )
    add_code_listing(
        doc,
        "4.17",
        "Próbkowanie dwuwymiarowego szumu Simplex z tablicą permutacji",
        SIMPLEX_LISTING,
        "opracowanie własne na podstawie [3] i [4]",
    )
    add_body(
        doc,
        "Fragment kodu zamieszczony na listingu 4.17 rozpoczyna się od transformacji skośnej i wyznaczenia "
        "indeksów i oraz j. Następnie obliczane są lokalne współrzędne pierwszego narożnika i wybierany jest "
        "drugi narożnik trójkąta. Wartości g0, g1 i g2 pochodzą z tablicy permutacji. Funkcja Kernel2DPerm "
        "realizuje równanie (4.3), a suma jej trzech wywołań jest skalowana przez współczynnik 70 i ograniczana "
        "do przedziału od -1 do 1. Kod odpowiada zatem kolejno: lokalizacji simpleksu, wyborowi gradientów, "
        "obliczeniu wkładów i normalizacji wyniku."
    )

    add_heading(doc, "4.1.4. Szum Worleya", 3)
    add_body(
        doc,
        "Szum Worleya, nazywany również szumem komórkowym, nie korzysta z gradientów. Przestrzeń zawiera "
        "pseudolosowo rozmieszczone punkty cech, a wartość funkcji jest wyznaczana na podstawie uporządkowanych "
        "odległości od punktu próbkowania [5]. Najczęściej używana jest odległość do najbliższego punktu F₁, "
        "lecz dostępne są również F₂ oraz kombinacje obu wartości."
    )
    add_body(
        doc,
        "Formalny zapis funkcji komórkowej wymaga uporządkowania odległości od wszystkich punktów cech. "
        "W praktyce sprawdzany jest ograniczony zbiór sąsiednich komórek, które mogą zawierać najbliższy punkt."
    )
    add_equation(
        doc,
        "4.4",
        "F_n(p) = n-ta najmniejsza wartość zbioru { d(p, q_i) }",
        "Definicja n-tej funkcji odległości szumu Worleya",
        [
            ("Fₙ(p)", "n-ta najmniejsza odległość dla punktu p"),
            ("p", "punkt próbkowania"),
            ("qᵢ", "i-ty punkt cechy"),
            ("d(p, qᵢ)", "wybrana metryka odległości między punktami"),
            ("n", "numer uporządkowanej odległości, na przykład 1 lub 2"),
        ],
        "S. Worley, „A Cellular Texture Basis Function”, SIGGRAPH 1996, DOI: 10.1145/237170.237267 [5].",
    )
    add_body(
        doc,
        "Równanie (4.4) wyjaśnia odmienny charakter obrazu Worleya widocznego na rysunku 4.1. Zmiana "
        "najbliższego punktu cechy tworzy granice komórek. W implementacji dwuwymiarowej przeszukiwane jest "
        "otoczenie 3 × 3 komórki zawierającej próbkę, a jednocześnie przechowywane są dwie najmniejsze "
        "odległości. Parametr jitter przesuwa punkt cechy względem środka komórki. Metryka euklidesowa, jej "
        "wersja bez pierwiastka oraz metryka Manhattan zmieniają geometrię uzyskiwanych obszarów."
    )

    add_heading(doc, "4.2. Funkcje fraktalne fBM i Ridged", 2)
    add_body(
        doc,
        "Pojedyncza funkcja szumu opisuje zmienność w ograniczonym paśmie częstotliwości, dlatego wygenerowany "
        "teren może być zbyt jednorodny i pozbawiony szczegółów. Synteza fraktalna łączy wiele próbek tej samej "
        "funkcji pobieranych w kolejnych skalach. Każda kolejna warstwa, nazywana oktawą, zwiększa częstotliwość "
        "zgodnie z lacunarity i zmniejsza amplitudę zgodnie z persistence [6, 7]."
    )
    add_heading(doc, "4.2.1. Fraktalny ruch Browna", 3)
    add_body(
        doc,
        "W grafice komputerowej fBM jest realizowany jako skończona suma ważonych oktaw funkcji bazowej. "
        "Normalizacja przez sumę amplitud stabilizuje zakres wyniku przy zmianie liczby oktaw, co ułatwia "
        "stosowanie wspólnej krzywej wysokości i porównywanie konfiguracji."
    )
    add_body(
        doc,
        "Częstotliwość i amplituda kolejnej oktawy tworzą ciągi geometryczne. Dla typowych wartości "
        "lacunarity równych 2 i persistence równych 0,5 następna oktawa zawiera dwukrotnie drobniejsze "
        "struktury o połowie amplitudy."
    )
    add_equation(
        doc,
        "4.5",
        "fBM(p) = [Σ(i=0,...,o-1) g^i N(λ^i p)] / [Σ(i=0,...,o-1) g^i]",
        "Znormalizowana suma oktaw funkcji szumu",
        [
            ("fBM(p)", "wartość fraktalnego ruchu Browna w punkcie p"),
            ("N(p)", "wartość funkcji szumu bazowego"),
            ("o", "liczba oktaw"),
            ("λ", "lacunarity, czyli mnożnik częstotliwości"),
            ("g", "persistence lub gain, czyli mnożnik amplitudy"),
            ("i", "indeks oktawy"),
        ],
        "F. K. Musgrave, C. E. Kolb, R. S. Mace, „The Synthesis and Rendering of Eroded Fractal Terrains”, SIGGRAPH 1989 [6]; D. S. Ebert i in., „Texturing and Modeling: A Procedural Approach”, 2003 [7].",
    )
    add_body(
        doc,
        "Równanie (4.5) odpowiada bezpośrednio parametrom octaves, lacunarity i persistence. Mnożnik λ "
        "skaluje współrzędne wejściowe, natomiast g skaluje amplitudę. Mianownik zapobiega wzrostowi "
        "maksymalnej wartości wraz z dodawaniem kolejnych oktaw."
    )
    add_body(
        doc,
        "Realizacja programowa przechowuje bieżącą sumę, amplitudę i sumę amplitud. Po każdej iteracji "
        "współrzędne są mnożone przez lacunarity, a amplituda przez gain."
    )
    add_code_listing(
        doc,
        "4.18",
        "Synteza fBM w dwóch wymiarach",
        FBM_LISTING,
        "opracowanie własne na podstawie [6] i [7]",
    )
    add_body(
        doc,
        "Na listingu 4.18 pętla wykonuje dokładnie liczbę iteracji określoną przez octaves. Każda próbka "
        "jest mnożona przez aktualną amplitudę i dodawana do sumy, po czym zmieniane są skala i ziarno "
        "kolejnej oktawy. Dzielenie przez maxAmp realizuje normalizację z równania (4.5)."
    )

    add_heading(doc, "4.2.2. Transformacja Ridged", 3)
    add_body(
        doc,
        "Transformacja grzbietowa zmienia sposób interpretacji wartości szumu. Zastosowanie wartości "
        "bezwzględnej składa ujemną i dodatnią część sygnału, a odwrócenie wyniku przekształca okolice "
        "zera w lokalne maksima. W efekcie zamiast łagodnych wzniesień powstają wydłużone grzbiety [7]."
    )
    add_body(
        doc,
        "W implementacji wynik pojedynczej oktawy jest dodatkowo przeskalowany z przedziału [0, 1] do "
        "przedziału [-1, 1]. Pozwala to używać tej samej procedury normalizacji i mieszania co dla fBM."
    )
    add_equation(
        doc,
        "4.6",
        "R(n) = 2(1 - |n|) - 1 = 1 - 2|n|",
        "Transformacja grzbietowa z normalizacją do przedziału [-1, 1]",
        [
            ("R(n)", "wartość próbki po transformacji grzbietowej"),
            ("n", "wartość szumu bazowego z przedziału [-1, 1]"),
            ("|n|", "wartość bezwzględna próbki"),
        ],
        "Transformacja grzbietowa na podstawie [7]; przekształcenie 2x − 1 jest normalizacją przedziału [0, 1] do [-1, 1].",
    )
    add_body(
        doc,
        "Równanie (4.6) przyjmuje wartość 1 dla n = 0 i wartość -1 dla n = -1 lub n = 1. Oznacza to, że "
        "obszary przejścia szumu przez zero stają się grzbietami, natomiast skrajne wartości stają się "
        "obniżeniami."
    )
    add_body(
        doc,
        "Wariant wielooktawowy stosuje tę samą progresję częstotliwości i amplitudy co fBM, ale przed "
        "zsumowaniem każdej próbki wykonuje transformację grzbietową."
    )
    add_code_listing(
        doc,
        "4.19",
        "Wielooktawowa transformacja Ridged w dwóch wymiarach",
        RIDGED_LISTING,
        "opracowanie własne na podstawie [7]",
    )
    add_body(
        doc,
        "Fragment kodu na listingu 4.19 realizuje równanie (4.6) w instrukcji zawierającej absN. "
        "Pozostała część pętli odpowiada syntezie fBM. Zastosowany wariant jest sumą oktaw po transformacji "
        "grzbietowej. Nie jest to pełny model ridged multifractal Musgrave'a, ponieważ amplituda kolejnej "
        "oktawy nie zależy od wyniku oktawy poprzedniej. Takie nazwanie wariantu zapobiega utożsamieniu "
        "dwóch różnych algorytmów."
    )
    add_body(
        doc,
        "Różnicę między jedną oktawą, fBM i transformacją grzbietową można ocenić przy wspólnej funkcji "
        "bazowej oraz tych samych parametrach skali."
    )
    add_figure(
        doc,
        figures["fractal"],
        "Rysunek 4.3. Porównanie jednej oktawy, fBM oraz wariantu Ridged",
        "Źródło: opracowanie własne na podstawie [6] i [7].",
    )
    add_body(
        doc,
        "Na rysunku 4.3 fBM dodaje coraz drobniejsze szczegóły, zachowując łagodną strukturę pierwszej "
        "oktawy. Wariant Ridged koncentruje jasne wartości w pobliżu miejsc, w których funkcja bazowa "
        "przechodzi przez zero. Zmiana jest skutkiem transformacji wartości, a nie użycia innej funkcji "
        "pseudolosowej."
    )

    add_heading(doc, "4.3. Łączenie różnych warstw szumu", 2)
    add_body(
        doc,
        "Pojedyncza konfiguracja fraktalna nadal opisuje jeden rodzaj struktury. Bardziej zróżnicowany "
        "teren można uzyskać przez połączenie kilku niezależnie skonfigurowanych warstw, na przykład "
        "łagodnej warstwy bazowej, pasma grzbietowego i maski komórkowej. Każda warstwa zwraca próbkę "
        "z własną częstotliwością, liczbą oktaw i siłą."
    )
    add_body(
        doc,
        "Łączenie warstw jest decyzją projektową aplikacji, a nie częścią definicji szumu Perlina, Simplex "
        "lub Worleya. Z tego względu operacje są wykonywane jawnie i w kolejności zapisanej w kolekcji. "
        "Pierwsza aktywna warstwa inicjuje wynik, a każda następna modyfikuje go zgodnie z wybranym trybem."
    )
    add_code_listing(doc, "4.20", "Sekwencyjne mieszanie aktywnych warstw szumu", BLEND_LISTING)
    add_body(
        doc,
        "Na listingu 4.20 próbka jest najpierw mnożona przez strength. Tryb Add dodaje ją do wyniku, "
        "Subtract odejmuje, natomiast Max i Min wybierają odpowiednio większą albo mniejszą wartość. "
        "W trybie Multiply ważona próbka jest przekształcana w nieujemny współczynnik mnożący. Po "
        "przetworzeniu wszystkich warstw wynik zostaje ograniczony do przedziału [-1, 1]. Kolejność ma "
        "znaczenie zwłaszcza dla odejmowania i mnożenia, dlatego konfiguracja powinna być traktowana jako "
        "uporządkowany stos, a nie nieuporządkowany zbiór."
    )

    add_heading(doc, "4.4. Generowanie fragmentu terenu", 2)
    add_body(
        doc,
        "Świat jest dzielony na regularne fragmenty generowane niezależnie. Każdy fragment zawiera siatkę "
        "wierzchołków o rozdzielczości zależnej od poziomu szczegółowości. Próbkowanie wykorzystuje "
        "współrzędne świata, dzięki czemu sąsiednie fragmenty odwołują się do tego samego ciągłego pola "
        "szumu i zachowują zgodność wysokości na wspólnych krawędziach."
    )
    add_body(
        doc,
        "Proces przekształca współrzędne w dane geometryczne etapami. Najpierw obliczana jest próbka "
        "stosu szumów, następnie kształtowana jest mapa wysokości, wyznaczany jest klimat, a na końcu "
        "zapisywane są atrybuty wierzchołków i indeksy trójkątów."
    )
    add_figure(
        doc,
        figures["pipeline"],
        "Rysunek 4.4. Przepływ danych podczas generowania fragmentu terenu",
        "Źródło: opracowanie własne.",
        width_cm=16.0,
    )
    add_body(
        doc,
        "Na rysunku 4.4 przedstawiono zależności między etapami. Warstwy szumu i funkcje fraktalne tworzą "
        "wartość bazową. Krzywa wysokości umożliwia nieliniowe przekształcenie rozkładu, na przykład "
        "spłaszczenie nizin lub podkreślenie szczytów. Dopiero ukształtowana wysokość jest używana do "
        "korekty temperatury, wilgotności i wyboru koloru biomu."
    )
    add_body(
        doc,
        "Aby uniknąć kosztownego wywoływania obiektu AnimationCurve w zadaniu, krzywa jest wcześniej "
        "próbkowana do tablicy LUT. Wartość szumu z przedziału [-1, 1] jest mapowana do [0, 1], a następnie "
        "używana jako indeks zmiennoprzecinkowy tablicy. Wartość pomiędzy sąsiednimi elementami LUT jest "
        "wyznaczana interpolacją liniową."
    )
    add_code_listing(
        doc,
        "4.21",
        "Wyznaczanie wysokości i danych klimatycznych punktu siatki",
        TERRAIN_EXECUTE_LISTING,
    )
    add_body(
        doc,
        "Fragment kodu zamieszczony na listingu 4.21 zamienia liniowy indeks zadania na współrzędne x i y "
        "siatki, a następnie na współrzędne świata. Próbka stosu warstw jest mapowana do dziedziny tablicy "
        "LUT i interpolowana między elementami i0 oraz i1. Po przemnożeniu przez amplitudę wynik zostaje "
        "ograniczony do poziomu morza. W tej samej iteracji wyznaczane są temperatura i wilgotność, a "
        "cztery wartości są zapisywane w strukturze float4, która stanowi wspólny bufor dalszych etapów."
    )
    add_body(
        doc,
        "Normalna wierzchołka wymaga wysokości po obu stronach punktu. Z tego powodu bufor wysokości "
        "zawiera dodatkową jednowierzchołkową ramkę. Pozwala ona obliczyć normalne na krawędziach fragmentu "
        "z użyciem tych samych różnic centralnych co w jego wnętrzu."
    )
    add_code_listing(doc, "4.22", "Obliczanie normalnej na podstawie sąsiednich wysokości", NORMAL_LISTING)
    add_body(
        doc,
        "Na listingu 4.22 odczytywane są wysokości lewego, prawego, dolnego i górnego sąsiada. Różnice "
        "hL − hR oraz hD − hU opisują nachylenie powierzchni, natomiast składowa pionowa uwzględnia "
        "odległość równą dwóm krokom siatki. Po normalizacji wektor jest zapisywany w buforze wierzchołków. "
        "Użycie ramki eliminuje potrzebę specjalnych przypadków dla brzegów i ogranicza różnice oświetlenia "
        "między fragmentami."
    )

    add_heading(doc, "4.5. Asynchroniczne wykonywanie generacji", 2)
    add_body(
        doc,
        "Generowanie siatki obejmuje dużą liczbę niezależnych próbek i dobrze nadaje się do przetwarzania "
        "równoległego. C# Job System rozdziela iteracje zadań IJobParallelFor między wątki robocze, a "
        "obiekty JobHandle opisują zależności i moment zakończenia pracy [11]. Kompilator Burst tłumaczy "
        "zgodny kod pośredni na zoptymalizowany kod natywny z użyciem LLVM [12]."
    )
    add_body(
        doc,
        "Kolejne zadania nie mogą jednak rozpocząć się w dowolnej kolejności. Normalne wymagają gotowych "
        "wysokości, a końcowy stan fragmentu powinien zostać zgłoszony dopiero po zakończeniu wszystkich "
        "obliczeń geometrii."
    )
    add_figure(
        doc,
        figures["jobs"],
        "Rysunek 4.5. Zależności pomiędzy zadaniami generowania siatki",
        "Źródło: opracowanie własne na podstawie dokumentacji Unity [11, 12].",
        width_cm=15.0,
    )
    add_body(
        doc,
        "Na rysunku 4.5 pierwsze zadanie wypełnia bufor wysokości, klimatu i kolorów. Drugie zadanie "
        "oblicza normalne i otrzymuje uchwyt pierwszego jako zależność. Zadanie budujące indeksy jest "
        "planowane po nim, a jego uchwyt reprezentuje zakończenie całego łańcucha. Menedżer okresowo "
        "sprawdza stan uchwytu bez blokowania głównego wątku."
    )
    add_body(
        doc,
        "Planowanie wykorzystuje partie po 64 iteracje. Wielkość partii ogranicza koszt dystrybucji pracy, "
        "a jednocześnie pozostawia środowisku wykonawczemu możliwość rozłożenia dużej siatki między wiele "
        "wątków."
    )
    add_code_listing(
        doc,
        "4.23",
        "Planowanie łańcucha zadań generowania fragmentu",
        SCHEDULE_LISTING,
        "opracowanie własne na podstawie dokumentacji Unity [11, 12]",
    )
    add_body(
        doc,
        "Na listingu 4.23 handle1 reprezentuje zakończenie próbkowania terenu. Przekazanie go do wywołania "
        "Schedule zadania normalnych wymusza poprawną kolejność bez ręcznego oczekiwania. Analogicznie "
        "handle2 staje się zależnością zadania indeksów. Dopiero handle3 jest zapisywany jako uchwyt całego "
        "fragmentu. Po wykryciu zakończenia wywoływana jest metoda Complete, a fragment trafia do kolejki "
        "aplikowania siatki. Sama modyfikacja obiektu Mesh pozostaje na wątku głównym, natomiast jej liczba "
        "i budżet czasu na klatkę są ograniczane przez menedżer terenu."
    )

    add_heading(doc, "4.6. Mapa temperatury i wilgotności", 2)
    add_body(
        doc,
        "Parametry klimatyczne są generowane jako dwa niezależne pola ciągłe próbkowane w tych samych "
        "współrzędnych świata co wysokość. Rozdzielenie źródeł pozwala uzyskać różne układy stref, ponieważ "
        "temperatura i wilgotność mogą używać odmiennych częstotliwości, ziaren oraz funkcji bazowych. "
        "Wartości końcowe są przechowywane w przedziale [0, 1]."
    )
    add_body(
        doc,
        "Po normalizacji surowych próbek stosowane są korekty zależne od wysokości. Temperatura maleje "
        "w kierunku szczytów, wilgotność jest zwiększana w dolinach i ograniczana na dużych wysokościach. "
        "Są to sterowalne reguły artystyczne generatora, a nie model fizyczny atmosfery, dlatego ich "
        "współczynniki pozostają parametrami konfiguracji."
    )
    add_code_listing(doc, "4.24", "Normalizacja i korekta parametrów klimatycznych", CLIMATE_LISTING)
    add_body(
        doc,
        "Na listingu 4.24 surowa temperatura jest najpierw przekształcana funkcją pierwiastkową zachowującą "
        "znak, co zwiększa udział wartości oddalonych od zera, a następnie mapowana do [0, 1]. Wilgotność "
        "korzysta z mapowania liniowego. Wysokość względna height01 steruje chłodzeniem i osuszaniem szczytów. "
        "Osobny współczynnik valleyWetness wzmacnia wilgotność poniżej ustalonej granicy dolin. Funkcja "
        "saturate po każdym etapie zabezpiecza zakres danych używanych przez klasyfikację biomów."
    )

    add_heading(doc, "4.7. Kolorowanie i rysowanie powierzchni terenu", 2)
    add_body(
        doc,
        "Kolor powierzchni jest zapisywany jako atrybut wierzchołka podczas generowania siatki. Dzięki temu "
        "wybór biomu nie wymaga osobnej tekstury mapy i pozostaje zgodny z wartościami wysokości, temperatury "
        "i wilgotności obliczonymi dla tego samego punktu. Kolorowanie łączy klasyfikację klimatyczną z "
        "pasmami zależnymi od wysokości."
    )
    add_body(
        doc,
        "Podstawowa paleta klimatyczna tworzy macierz trzech poziomów temperatury i trzech poziomów "
        "wilgotności. Wartości pośrednie są interpolowane, dlatego przejścia między biomami nie tworzą "
        "ostrych granic."
    )
    add_figure(
        doc,
        figures["biomes"],
        "Rysunek 4.6. Macierz bazowych kolorów biomów",
        "Źródło: opracowanie własne na podstawie palety zastosowanej w aplikacji.",
        width_cm=14.5,
    )
    add_body(
        doc,
        "Na rysunku 4.6 przedstawiono dziewięć punktów odniesienia palety. Przesuwanie się w prawo zwiększa "
        "wilgotność i prowadzi od barw suchych do zieleni. Przesuwanie się w dół zwiększa temperaturę, "
        "zmieniając barwy chłodne w żółte i pomarańczowe. Funkcja wyboru strefy interpoluje najpierw w "
        "jednym, a następnie w drugim wymiarze macierzy."
    )
    add_body(
        doc,
        "Wysokość ma pierwszeństwo przed klasyfikacją klimatu w obszarach morza i plaży. Powyżej strefy "
        "biomu do koloru dodawana jest skała, a następnie śnieg. Granica śniegu zależy od temperatury, "
        "dzięki czemu w zimnych regionach pojawia się na mniejszej wysokości."
    )
    add_code_listing(doc, "4.25", "Wyznaczanie koloru powierzchni terenu", COLOR_LISTING)
    add_body(
        doc,
        "Fragment kodu zamieszczony na listingu 4.25 najpierw obsługuje morze i pas plaży. Następnie "
        "GetZoneColor wyznacza kolor klimatyczny na podstawie temperatury oraz wilgotności. Funkcja "
        "smoothstep tworzy łagodne przejście z plaży do biomu i z biomu do skały. Ostatni etap oblicza "
        "temperaturowo zależną granicę śniegu i interpoluje kolor w paśmie snowBandWidth. Otrzymany Color32 "
        "jest zapisywany razem z pozycją i normalną w buforze wierzchołka, a po zakończeniu zadań przekazywany "
        "do siatki renderowanej przez silnik."
    )

    add_heading(doc, "Bibliografia wykorzystana w rozdziale", 2)
    bibliography = [
        "[1] K. Perlin, „An Image Synthesizer”, ACM SIGGRAPH Computer Graphics, 19(3), 287–296, 1985. DOI: 10.1145/325165.325247.",
        "[2] K. Perlin, „Improving Noise”, Proceedings of SIGGRAPH 2002, 681–682, 2002. DOI: 10.1145/566570.566636.",
        "[3] K. Perlin, „Noise Hardware”, Real-Time Shading, SIGGRAPH Course Notes, 2001.",
        "[4] S. Gustavson, „Simplex noise demystified”, Linköping University, 2005. https://itn-web.it.liu.se/~stegu76/TNM084-2011/simplexnoise-demystified.pdf",
        "[5] S. Worley, „A Cellular Texture Basis Function”, Proceedings of SIGGRAPH 1996, 291–294. DOI: 10.1145/237170.237267.",
        "[6] F. K. Musgrave, C. E. Kolb, R. S. Mace, „The Synthesis and Rendering of Eroded Fractal Terrains”, Proceedings of SIGGRAPH 1989, 41–50. DOI: 10.1145/74333.74337.",
        "[7] D. S. Ebert, F. K. Musgrave, D. Peachey, K. Perlin, S. Worley, „Texturing and Modeling: A Procedural Approach”, 3rd ed., Morgan Kaufmann, 2003.",
        "[8] A. Lagae, P. Dutré, „Long-Period Hash Functions for Procedural Texturing”, Vision, Modeling, and Visualization 2006, 225–228. https://graphics.cs.kuleuven.be/publications/LD06LPHFPT/",
        "[9] R. A. Fisher, F. Yates, „Statistical Tables for Biological, Agricultural and Medical Research”, Oliver & Boyd, 1938.",
        "[10] G. Marsaglia, „Xorshift RNGs”, Journal of Statistical Software, 8(14), 1–6, 2003. DOI: 10.18637/jss.v008.i14.",
        "[11] Unity Technologies, „JobHandle i zależności”, dokumentacja C# Job System. https://docs.unity3d.com/Manual/JobSystemJobDependencies.html",
        "[12] Unity Technologies, „Burst”, dokumentacja pakietu kompilatora Burst. https://docs.unity3d.com/Manual/com.unity.burst.html",
    ]
    for entry in bibliography:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(entry)
        set_run_font(r, "Times New Roman", 10.5)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


def audit_forbidden_terms(path):
    forbidden = ("FastNoiseLite", "RewriteFastNoiseLite")
    with zipfile.ZipFile(path) as zf:
        xml = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.endswith(".xml")
        )
    found = [term for term in forbidden if term in xml]
    if found:
        raise RuntimeError(f"Forbidden terms found in document: {found}")


def main():
    figures = generate_figures()
    output = build_document(figures)
    audit_forbidden_terms(output)
    print(output)


if __name__ == "__main__":
    main()
